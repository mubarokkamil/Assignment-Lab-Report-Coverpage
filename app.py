from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

# Function to replace text while preserving formatting
def replace_text_preserve_format(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            full_text = "".join(run.text for run in paragraph.runs)
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text

# Replace placeholders in docx and convert to PDF using LibreOffice
def replace_placeholders(input_file, output_file, pdf_file, replacements):
    doc = Document(input_file)

    for para in doc.paragraphs:
        replace_text_preserve_format(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_preserve_format(para, replacements)

    doc.save(output_file)

    # Convert to PDF using LibreOffice
    os.system(f'libreoffice --headless --convert-to pdf "{output_file}" --outdir "{os.getcwd()}"')
    print(f"Converted {output_file} to PDF at {pdf_file}.")

@app.route("/", methods=["GET", "POST"])
def choose_type():
    if request.method == "POST":
        choice = request.form.get("choice")
        if choice == "cover":
            return render_template("cover_form.html")
        elif choice == "lab":
            return render_template("lab_form.html")
    return render_template("choice.html")

@app.route("/generate_cover", methods=["POST"])
def generate_cover():
    replacements = {
        "{Course_Code}": request.form["course_code"],
        "{Course_Name}": request.form["course_name"],
        "{Course_Teacher}": request.form["course_teacher"],
        "{Course_Teacher_Details}": request.form["course_teacher_details"],
        "{Assignment_Name}": request.form["assignment_name"],
        "{Reg_No}": request.form["reg_no"],
        "{Name}": request.form["name"]
    }

    input_file = os.path.join(os.getcwd(), "CoverPage.docx")
    output_file = os.path.join(os.getcwd(), "Modified_CoverPage.docx")
    pdf_file = os.path.join(os.getcwd(), "Modified_CoverPage.pdf")

    if not os.path.exists(input_file):
        return "Error: CoverPage.docx not found!", 400

    replace_placeholders(input_file, output_file, pdf_file, replacements)

    return send_file(pdf_file, as_attachment=True)

@app.route("/generate_lab", methods=["POST"])
def generate_lab():
    replacements = {
        "{Course_Code}": request.form["course_code"],
        "{Course_Name}": request.form["course_name"],
        "{Course_Teacher}": request.form["course_teacher"],
        "{Course_Teacher_Details}": request.form["course_teacher_details"],
        "{Assignment_Name}": request.form["assignment_name"],
        "{Reg_No}": request.form["reg_no"],
        "{Name}": request.form["name"]
    }

    input_file = os.path.join(os.getcwd(), "Lab_Report.docx")
    output_file = os.path.join(os.getcwd(), "Modified_Lab_Report.docx")
    pdf_file = os.path.join(os.getcwd(), "Modified_Lab_Report.pdf")

    if not os.path.exists(input_file):
        return "Error: Lab_Report.docx not found!", 400

    replace_placeholders(input_file, output_file, pdf_file, replacements)

    return send_file(pdf_file, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
